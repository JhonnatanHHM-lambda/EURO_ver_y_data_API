"""
Generador de Word (.docx) del Acta de Carga de Datos.
Usa PLANTILLA LOGO CON OPACIDAD.png como fondo de página (mismo patrón
que euro_generador_actas/core/utils/word_generator.py).
"""
import base64
import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement

from .acta_brand import PLANTILLA_PNG, LOGO_EURO

# ── Márgenes idénticos al PDF ──────────────────────────────────────────────────
_TOP   = Pt(115)
_BOT   = Pt(68)
_LEFT  = Pt(54)
_RIGHT = Pt(54)

# ── Colores ────────────────────────────────────────────────────────────────────
_AZUL  = RGBColor(0x27, 0x34, 0x8B)
_VERDE = RGBColor(0x16, 0xA3, 0x4A)
_ROJO  = RGBColor(0xDC, 0x26, 0x26)
_TEXTO = RGBColor(0x2C, 0x2C, 0x2C)
_SUAVE = RGBColor(0x66, 0x66, 0x66)
_BLANCO= RGBColor(0xFF, 0xFF, 0xFF)
_BORDE = "D8DCF0"

# Dimensiones página carta en EMU
_W_EMU   = 7772400
_H_EMU   = 10058400
_PAGE_W  = 8.5 - 0.75 - 0.75   # ancho de contenido en pulgadas


# ── Fondo de página (igual que word_generator.py) ─────────────────────────────

def _add_page_bg(header, img_path):
    p   = header.paragraphs[0]
    run = p.add_run()
    run.add_picture(img_path, width=Inches(8.5), height=Inches(11))

    drawing_el = run._r.find(qn('w:drawing'))
    inline_el  = drawing_el.find(qn('wp:inline'))
    blip_el    = inline_el.find('.//' + qn('a:blip'))
    r_id       = blip_el.get(qn('r:embed'))
    doc_pr     = inline_el.find(qn('wp:docPr'))
    uid        = doc_pr.get('id', '1')
    name       = doc_pr.get('name', 'bg')
    run._r.remove(drawing_el)

    anchor_xml = (
        f'<w:drawing'
        f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        f' xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        f' simplePos="0" relativeHeight="251657216" behindDoc="1"'
        f' locked="1" layoutInCell="1" allowOverlap="0">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{_W_EMU}" cy="{_H_EMU}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{uid}" name="{name}"/>'
        f'<wp:cNvGraphicFramePr>'
        f'<a:graphicFrameLocks noChangeAspect="1"/>'
        f'</wp:cNvGraphicFramePr>'
        f'<a:graphic>'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic>'
        f'<pic:nvPicPr>'
        f'<pic:cNvPr id="{uid}" name="{name}"/>'
        f'<pic:cNvPicPr><a:picLocks noChangeAspect="1" noChangeArrowheads="1"/></pic:cNvPicPr>'
        f'</pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip r:embed="{r_id}"/>'
        f'<a:stretch><a:fillRect/></a:stretch>'
        f'</pic:blipFill>'
        f'<pic:spPr bwMode="auto">'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{_W_EMU}" cy="{_H_EMU}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:noFill/>'
        f'</pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
    )
    run._r.append(parse_xml(anchor_xml))


# ── Helpers de párrafo ─────────────────────────────────────────────────────────

def _para(doc, text='', bold=False, size=10, color=None,
          space_before=0, space_after=4, align='justify'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = {
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        'center':  WD_ALIGN_PARAGRAPH.CENTER,
        'left':    WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color or _TEXTO
    return p


def _para_inline(doc, items, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold, size, color in items:
        r = p.add_run(text)
        r.bold = bold
        r.font.size  = Pt(size)
        r.font.color.rgb = color or _TEXTO
    return p


# ── Helpers de tabla ───────────────────────────────────────────────────────────

def _tbl_borders(tbl):
    tbl._tbl.tblPr.append(parse_xml(
        f'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top    w:val="single" w:sz="6" w:color="{_BORDE}"/>'
        f'<w:left   w:val="single" w:sz="6" w:color="{_BORDE}"/>'
        f'<w:bottom w:val="single" w:sz="6" w:color="{_BORDE}"/>'
        f'<w:right  w:val="single" w:sz="6" w:color="{_BORDE}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:color="{_BORDE}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:color="{_BORDE}"/>'
        f'</w:tblBorders>'
    ))


def _cell_width(cell, width_in):
    tc_pr = cell._tc.get_or_add_tcPr()
    for el in tc_pr.findall(qn('w:tcW')):
        tc_pr.remove(el)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_in * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tc_pr.append(tcW)


def _cell_fill(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    ))


def _cell_pad(cell, pt=6):
    dxa = int(pt * 20 * 72 / 72)
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:w="80" w:type="dxa"/>'
        f'<w:left w:w="120" w:type="dxa"/>'
        f'<w:bottom w:w="80" w:type="dxa"/>'
        f'<w:right w:w="120" w:type="dxa"/>'
        f'</w:tcMar>'
    ))


def _cell_vcenter(cell):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:val="center"/>'
    ))


def _cell_run(cell, text, bold=False, size=10, color=None, align='left', italic=False):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if align == 'left' else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or _TEXTO
    return r


def _yellow_header_line(tbl):
    """Agrega línea amarilla debajo de la primera fila (igual que en PDF)."""
    tbl._tbl.tblPr.find(qn('w:tblBorders')).append(parse_xml(
        f'<w:insideH xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:val="single" w:sz="4" w:color="{_BORDE}"/>'
    ))


# ── Tabla información de la carga ─────────────────────────────────────────────

def _tabla_info(doc, carga):
    fecha_carga = carga.creado.strftime('%d/%m/%Y %H:%M') if carga.creado else '—'
    cargado_por = (carga.cargado_por.obtener_nombre_completo()
                   if carga.cargado_por else '—')
    sede_nombre = carga.sede.nombre if carga.sede else 'Sin sede asignada'

    c1 = _PAGE_W * 0.33
    c2 = _PAGE_W * 0.33
    c3 = _PAGE_W * 0.34

    tbl = doc.add_table(rows=4, cols=3)
    _tbl_borders(tbl)

    rows_data = [
        [('ARCHIVO', True), ('HOJA', True), ('ORIGEN', True)],
        [(carga.nombre_archivo, False), (carga.hoja or '—', False), (carga.origen_datos, False)],
        [('SEDE', True), ('CARGADO POR', True), ('FECHA DE CARGA', True)],
        [(sede_nombre, False), (cargado_por, False), (fecha_carga, False)],
    ]
    hdr_rows = {0, 2}
    for ri, row_data in enumerate(rows_data):
        row = tbl.rows[ri]
        is_hdr = ri in hdr_rows
        for ci, (text, bold) in enumerate(row_data):
            cell = row.cells[ci]
            _cell_width(cell, [c1, c2, c3][ci])
            _cell_pad(cell)
            if is_hdr:
                _cell_fill(cell, '27348B')
            else:
                _cell_fill(cell, 'F5F6FA')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(9 if is_hdr else 9.5)
            r.font.color.rgb = _BLANCO if is_hdr else _TEXTO
    return tbl


# ── Tabla estadísticas (KPIs) ─────────────────────────────────────────────────

def _tabla_kpis(doc, carga):
    col = _PAGE_W / 3
    labels = ['TOTAL EN ARCHIVO', 'CARGADOS', 'CON ERROR']
    values = [str(carga.total_registros), str(carga.exitosos), str(carga.fallidos)]
    colors = [_AZUL, _VERDE, _ROJO if carga.fallidos > 0 else _SUAVE]
    bgs    = ['F5F6FA', 'F0FDF4', 'FFF5F5' if carga.fallidos > 0 else 'F5F6FA']

    tbl = doc.add_table(rows=2, cols=3)
    _tbl_borders(tbl)

    for ci in range(3):
        # Cabecera
        hc = tbl.rows[0].cells[ci]
        _cell_width(hc, col); _cell_fill(hc, '27348B'); _cell_pad(hc)
        hp = hc.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(5)
        hp.paragraph_format.space_after  = Pt(5)
        hr = hp.add_run(labels[ci])
        hr.bold = True; hr.font.size = Pt(8); hr.font.color.rgb = _BLANCO

        # Valor
        vc = tbl.rows[1].cells[ci]
        _cell_width(vc, col); _cell_fill(vc, bgs[ci]); _cell_pad(vc)
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vp.paragraph_format.space_before = Pt(8)
        vp.paragraph_format.space_after  = Pt(8)
        vr = vp.add_run(values[ci])
        vr.bold = True; vr.font.size = Pt(18); vr.font.color.rgb = colors[ci]
    return tbl


# ── Tabla errores ─────────────────────────────────────────────────────────────

def _tabla_errores(doc, errores):
    LABELS = {
        'sin_documento':      'Sin número de documento',
        'sin_nombre':         'Sin nombre completo',
        'duplicado':          'Registro duplicado',
        'error_inesperado':   'Error de formato / longitud',
        'nombre_incoherente': 'Nombre incoherente con BD',
    }
    grupos = {}
    for e in errores:
        t = e.get('tipo', 'error_inesperado')
        grupos.setdefault(t, []).append(e.get('fila', '?'))

    c1 = _PAGE_W * 0.45; c2 = _PAGE_W * 0.1; c3 = _PAGE_W * 0.45
    tbl = doc.add_table(rows=1 + len(grupos), cols=3)
    _tbl_borders(tbl)

    for ci, (text, w) in enumerate(zip(['TIPO DE ERROR', 'CANT.', 'FILAS'], [c1, c2, c3])):
        hc = tbl.rows[0].cells[ci]
        _cell_width(hc, w); _cell_fill(hc, 'DC2626'); _cell_pad(hc)
        hp = hc.paragraphs[0]
        hp.paragraph_format.space_before = Pt(4); hp.paragraph_format.space_after = Pt(4)
        hr = hp.add_run(text)
        hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = _BLANCO

    for ri, (tipo, filas) in enumerate(grupos.items(), 1):
        filas_txt = ', '.join(str(f) for f in filas[:12])
        if len(filas) > 12: filas_txt += f' (+{len(filas)-12} más)'
        for ci, (text, w) in enumerate(zip(
                [LABELS.get(tipo, tipo), str(len(filas)), filas_txt], [c1, c2, c3])):
            cell = tbl.rows[ri].cells[ci]
            _cell_width(cell, w); _cell_pad(cell)
            if ri % 2 == 0: _cell_fill(cell, 'F5F6FA')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text); r.font.size = Pt(9); r.font.color.rgb = _TEXTO
    return tbl


# ── Tabla de firma (1 columna: Gestión Humana) ────────────────────────────────

def _tabla_firmas(doc, carga):
    tbl = doc.add_table(rows=1, cols=1)
    _tbl_borders(tbl)

    cell = tbl.rows[0].cells[0]
    _cell_width(cell, _PAGE_W); _cell_fill(cell, 'EEF0F8'); _cell_pad(cell)

    # Encabezado
    gp = cell.paragraphs[0]
    gp.paragraph_format.space_before = Pt(6); gp.paragraph_format.space_after = Pt(8)
    gr = gp.add_run('GESTIÓN HUMANA · EURO SUPERMERCADOS')
    gr.bold = True; gr.font.size = Pt(8); gr.font.color.rgb = _AZUL

    # Firma dibujada o línea en blanco
    if carga.firma_gh_imagen:
        try:
            img_data = carga.firma_gh_imagen
            if img_data.startswith('data:'):
                img_data = img_data.split(',', 1)[1]
            img_bytes = base64.b64decode(img_data)
            sp = cell.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after  = Pt(4)
            sp.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(1.8))
        except Exception:
            sp = cell.add_paragraph('________________________________')
            sp.paragraph_format.space_after = Pt(20)
            if sp.runs: sp.runs[0].font.size = Pt(11); sp.runs[0].font.color.rgb = _AZUL
    else:
        sp = cell.add_paragraph('________________________________')
        sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(20)
        if sp.runs: sp.runs[0].font.size = Pt(11); sp.runs[0].font.color.rgb = _AZUL

    nombre = carga.firma_gh_nombre or '___________________'
    cargo  = carga.firma_gh_cargo  or '___________________'
    fecha  = (carga.firma_gh_fecha.strftime('%d/%m/%Y %H:%M')
              if carga.firma_gh_fecha else '___/___/_______')

    for text, bold, size, color in [
        (nombre,          True,  9, _TEXTO),
        (cargo,           False, 8, _SUAVE),
        (f'Fecha: {fecha}', False, 8, _SUAVE),
    ]:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
        r.font.color.rgb = color

    return tbl


# ── Función principal ─────────────────────────────────────────────────────────

def generar_word_acta(carga) -> bytes:
    """Genera el .docx del acta y retorna bytes."""
    doc  = Document()
    sect = doc.sections[0]
    sect.top_margin      = _TOP
    sect.bottom_margin   = _BOT
    sect.left_margin     = _LEFT
    sect.right_margin    = _RIGHT
    sect.header_distance = Pt(0)
    sect.footer_distance = Pt(0)

    # Fuente por defecto
    doc.styles['Normal'].font.name      = 'Montserrat'
    doc.styles['Normal'].font.size      = Pt(10)
    doc.styles['Normal'].font.color.rgb = _TEXTO
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after  = Pt(0)
    doc.styles['Normal'].paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Fondo de página (membrete)
    if PLANTILLA_PNG:
        try:
            _add_page_bg(sect.header, PLANTILLA_PNG)
        except Exception:
            pass

    # Footer vacío
    ftr_p = sect.footer.paragraphs[0]
    ftr_p.paragraph_format.space_before = Pt(0)
    ftr_p.paragraph_format.space_after  = Pt(0)

    revertida = not carga.estado

    # ── Banner REVERTIDA ──────────────────────────────────────────────────
    if revertida:
        fecha_rev = carga.modificado.strftime('%d/%m/%Y %H:%M') if carga.modificado else '—'
        bt = doc.add_table(1, 1)
        _tbl_borders(bt)
        bc = bt.rows[0].cells[0]
        _cell_fill(bc, 'FEE2E2'); _cell_pad(bc)
        bp = bc.paragraphs[0]
        bp.paragraph_format.space_before = Pt(6); bp.paragraph_format.space_after = Pt(6)
        br = bp.add_run(f'⚠  CARGA REVERTIDA  —  Registros eliminados el {fecha_rev}.'
                        f' Este documento es el acta histórica del proceso.')
        br.bold = True; br.font.size = Pt(9)
        br.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
        _para(doc, space_after=8)

    # ── Título ────────────────────────────────────────────────────────────
    _para(doc, 'ACTA DE CARGA DE DATOS', bold=True, size=12, space_after=6)

    # ── 1. Información ────────────────────────────────────────────────────
    _para(doc, '1. Información de la carga', bold=True, size=11,
          space_before=10, space_after=5)
    _tabla_info(doc, carga)
    _para(doc, space_after=8)

    # ── 2. Estadísticas ───────────────────────────────────────────────────
    _para(doc, '2. Estadísticas de homologación', bold=True, size=11,
          space_before=10, space_after=5)
    _tabla_kpis(doc, carga)
    _para(doc, space_after=8)

    # ── 3. Errores ────────────────────────────────────────────────────────
    errores = carga.errores or []
    _para(doc, '3. Detalle de registros con error', bold=True, size=11,
          space_before=10, space_after=5)
    if errores:
        _tabla_errores(doc, errores)
    else:
        p = doc.add_paragraph('No se registraron errores. Todos los registros '
                               'fueron procesados correctamente.')
        p.paragraph_format.space_after = Pt(4)
        if p.runs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = _VERDE
    _para(doc, space_after=8)

    # ── 4. Certificación técnica ──────────────────────────────────────────
    _para(doc, '4. Certificación técnica', bold=True, size=11,
          space_before=10, space_after=5)
    cert = doc.add_paragraph(
        'El presente proceso de carga y homologación fue ejecutado siguiendo '
        'los estándares de calidad definidos en el Acta de Inicio del proyecto. '
        'Los datos han sido homologados al modelo unificado de trazabilidad de '
        'Euro Supermercados, con validación de campos obligatorios, detección '
        'de duplicados, inferencia de estados y limpieza de valores incorrectos. '
        'Este documento constituye evidencia técnica del proceso ejecutado.'
    )
    cert.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cert.paragraph_format.space_after = Pt(4)
    if cert.runs:
        cert.runs[0].font.size = Pt(10)
        cert.runs[0].font.color.rgb = _TEXTO

    _para(doc, space_after=20)

    # ── 5. Firmas ─────────────────────────────────────────────────────────
    _para(doc, '5. Firmas', bold=True, size=11, space_before=10, space_after=5)
    _tabla_firmas(doc, carga)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
